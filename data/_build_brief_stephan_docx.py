"""Brief Stephan en Word (.docx) — tables propres, pas de chevauchement."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "Brief_Stephan_Newsletters_Juin_2026.docx"

GREEN = RGBColor(0x0C, 0x4A, 0x30)
BROWN = RGBColor(0x6B, 0x44, 0x23)
TAN   = RGBColor(0xD4, 0xA3, 0x73)
CREAM_HEX = "F4E9D8"
GREEN_HEX = "0C4A30"

doc = Document()

# Set landscape A4
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(20)
        r.font.color.rgb = GREEN
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), GREEN_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        r.font.size = Pt(14)
        r.font.color.rgb = GREEN
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = BROWN
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(10)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10)

def add_table(headers, rows, col_widths_cm=None):
    """Add styled table — headers green, alternating cream rows."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.allow_autofit = False

    # Set fixed layout
    tblPr = tbl._tbl.tblPr
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Column widths
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)

    # Headers
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(cell, GREEN_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            # Detect "**bold**" markup
            if val.startswith("**") and val.endswith("**"):
                run = para.add_run(val[2:-2])
                run.bold = True
            else:
                run = para.add_run(val)
            run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r_idx % 2 == 1:
                shade_cell(cell, CREAM_HEX)
    return tbl

def add_hr():
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D4A373")
    pBdr.append(bottom)
    p_pr.append(pBdr)

# ===== CONTENT =====

add_heading("Brief Stephan — Newsletters Juin 2026", 1)

p = doc.add_paragraph()
p.add_run("À : ").bold = True
p.add_run("Stephan    ")
p.add_run("De : ").bold = True
p.add_run("Nicolas Raes    ")
p.add_run("Date : ").bold = True
p.add_run("11 mai 2026    ")
p.add_run("Envoi prévu : ").bold = True
p.add_run("mardi 9 juin 2026, 9h00")

add_hr()

# ----- Contexte -----
add_heading("Le contexte en 30 secondes", 2)
add_para(
    "On démarre une roadmap B2B sur 12 mois (juin 2026 → mai 2027) avec un objectif simple : "
    "+200 K€ de CA additionnel sur la base actuelle de 881 K€. La mécanique : 1 ou 2 campagnes Mailchimp "
    "ciblées par mois, segment par segment (Horeca, Revendeur, GMS, Grossiste), avec une offre claire chacune."
)
add_para(
    "Tu pilotes la partie newsletter Mailchimp de bout en bout. Moi je m'occupe des données Odoo "
    "(segmentation, tags, tracking). Vanessa pilote la suite : commandes entrantes, application de l'offre, "
    "SAV. Jérôme appuie en relance commerciale J+10 sur les comptes prioritaires."
)
p = doc.add_paragraph()
p.add_run("Principe juin : ").bold = True
p.add_run(
    "on cible UNIQUEMENT les clients dormants (qui n'ont pas commandé depuis > 6 mois). "
    "On n'envoie PAS aux Horeca actifs — pas de promo à des clients qui commandent déjà = pas de cannibalisation."
)
add_hr()

# ----- Tableau synthèse -----
add_heading("Juin 2026 — 2 campagnes en parallèle", 2)
add_table(
    headers=["Campagne", "Volume", "Offre", "CA cible"],
    rows=[
        ["**A — Horeca dormant** (sans commande > 6 mois)", "11",  "3+1 sur boîtes HC25 + relance Jérôme top 5", "+1 K€"],
        ["**B — Revendeur dormant** (sans commande > 6 mois)", "105", "-10% sur réassort libre ≥ 300€ + relance Jérôme top 20", "+8 K€"],
        ["**TOTAL**", "**116**", "", "**+9 K€**"],
    ],
    col_widths_cm=[10, 2.5, 11, 2.5]
)

add_para("Tags Odoo correspondants (à utiliser pour créer les 2 audiences Mailchimp) :", bold=True)
add_bullet("A → MC-202606-HORECA-DORMANT")
add_bullet("B → MC-202606-REVENDEUR-DORMANT")
add_hr()

# ----- Tracking sans code -----
add_heading("Tracking de la conversion (sans code promo)", 2)
add_para(
    "On ne demande PAS aux clients de saisir un code. C'est trop compliqué à suivre côté opérationnel.",
    bold=True
)
add_para("À la place, on tracke automatiquement via croisement de listes Odoo :")
add_bullet("On sait qui a reçu chaque mailing (tag Odoo MC-202606-HORECA-DORMANT ou MC-202606-REVENDEUR-DORMANT)")
add_bullet("On sait qui a commandé (sale.order date_order entre le 9 juin et le 9 juillet)")
add_bullet("Croisement → conversions identifiées automatiquement, sans code à saisir")
p = doc.add_paragraph()
p.add_run("Workflow Vanessa : ").bold = True
p.add_run(
    "elle traite les commandes comme d'habitude. Si l'offre est mentionnée par le client (3+1 ou -10%), "
    "elle l'applique sur le devis. Pas de code à saisir, pas de champ supplémentaire."
)
p = doc.add_paragraph()
p.add_run("Reporting J+30 (9 juillet) : ").bold = True
p.add_run(
    "je sors la liste \"partners taggés campagne juin qui ont passé commande entre 9 juin et 9 juillet\" "
    "→ liste exhaustive des conversions, CA généré, taux de réactivation."
)
add_hr()

# ----- Template A -----
add_heading("Les 2 templates à concevoir", 2)
add_heading("Template A — Horeca dormant (11 destinataires)", 3)

p = doc.add_paragraph()
p.add_run("Ton : ").bold = True
p.add_run("\"Vous nous manquez\", retour aux fondamentaux, sincère mais pas larmoyant")
p = doc.add_paragraph()
p.add_run("Objet email proposé : ").bold = True
p.add_run("« Cela fait un moment... 3+1 pour reprendre tranquillement ? »")
p = doc.add_paragraph()
p.add_run("Angle : ").bold = True
p.add_run("On a remarqué qu'on ne s'est plus parlé. Voici une offre pour reprendre contact sans pression.")
p = doc.add_paragraph()
p.add_run("Particularité : ").bold = True
p.add_run("Jérôme va appeler le top 5 de cette liste à J+10 si pas de retour")

add_para("Contenu obligatoire :", bold=True)
add_bullet("Offre : 3 boîtes HC25 achetées = 4ᵉ offerte, 1 réf au choix par tranche de 3")
add_bullet("Min 6 boîtes, port offert dès 200 €, valable juin uniquement")
add_bullet("Phrase d'accroche : « Nous avons remarqué que cela fait quelques mois que nous ne nous sommes pas parlés... »")
add_bullet("CTA simple : « Répondez à cet email pour passer commande, ou appelez Vanessa »")
add_bullet("Mention « Jérôme reste à votre disposition pour un échange, n'hésitez pas »")

# ----- Template B -----
add_heading("Template B — Revendeur dormant (105 destinataires)", 3)

p = doc.add_paragraph()
p.add_run("Ton : ").bold = True
p.add_run("Libre, respectueux, « vous choisissez ce que vous voulez », pas une opé promo classique")
p = doc.add_paragraph()
p.add_run("Objet email proposé : ").bold = True
p.add_run("« Reprenons contact — votre choix, -10% sur toute la gamme 🍵 »")
p = doc.add_paragraph()
p.add_run("Angle : ").bold = True
p.add_run("Pas de SKU forcé, pas de pack imposé. C'est votre réassort, à votre rythme. Juste -10% pour reprendre tranquillement.")
p = doc.add_paragraph()
p.add_run("Particularité : ").bold = True
p.add_run("Jérôme va appeler le top 20 (CA historique ≥ 2 K€) à J+10")

add_para("Contenu obligatoire :", bold=True)
add_bullet("Offre : -10% sur toute commande ≥ 300 €, frais de port offerts dès 250 €")
add_bullet("1 commande maximum par client, valable juin uniquement")
add_bullet("Phrase d'accroche : « Cela fait un moment que nous n'avons plus eu de vos nouvelles. Plutôt que vous proposer un assortiment imposé, nous vous laissons carte blanche... »")
add_bullet("CTA simple : « Passez commande via notre catalogue ou contactez Vanessa »")
add_bullet("Lien vers le catalogue produits (page Shopify pro)")
add_hr()

# ----- Calendrier -----
add_heading("Calendrier de production", 2)
add_table(
    headers=["Date", "Échéance", "Qui"],
    rows=[
        ["**11 mai (fait)**",       "Roadmap validée + tags Odoo posés + CSV généré",                            "Nicolas"],
        ["15 mai (J-25)",           "Brief créatif validé (objets, angles, visuels)",                            "Stephan + Nicolas"],
        ["20 mai (J-20)",           "Import CSV dans Mailchimp + création des 2 audiences",                      "Stephan"],
        ["25 mai (J-15)",           "Premiers drafts des 2 templates Mailchimp",                                 "Stephan"],
        ["27 mai (J-13)",           "Revue templates + ajustements",                                             "Nicolas"],
        ["2 juin (J-7)",            "Envoi de test interne (nicolas, vanessa, stephan, jerome)",                 "Stephan"],
        ["3 juin (J-6)",            "Brief Jérôme : top 20 Revendeur + top 5 Horeca dormants",                   "Nicolas"],
        ["**9 juin 9h00 (J)**",     "**Envoi des 2 campagnes (par segment)**",                                   "Stephan"],
        ["12 juin (J+3)",           "Relance Mailchimp aux non-ouvreurs (objet alternatif)",                     "Stephan"],
        ["19 juin (J+10)",          "Appels relance Jérôme top 20 + top 5",                                      "Jérôme"],
        ["9 juillet (J+30)",        "Reporting CA + ROI campagne via croisement listes Odoo",                    "Nicolas"],
    ],
    col_widths_cm=[4, 16, 6]
)
add_hr()

# ----- KPI Mailchimp -----
add_heading("KPI à tracker", 2)
add_heading("Métriques Mailchimp (que tu nous remontes)", 3)
add_table(
    headers=["Métrique", "Cible"],
    rows=[
        ["Taux d'ouverture moyen sur les 2 campagnes", "≥ 35% (sectoriel B2B premium) — visuel \"vous nous manquez\" booste l'ouverture"],
        ["Taux de clic moyen",                          "≥ 8%"],
        ["Désabonnements",                              "< 0,5%"],
        ["Bounces",                                     "5-10 sur 116 destinataires (~5-8%) — emails dormants = base partiellement périmée"],
    ],
    col_widths_cm=[10, 16]
)

add_heading("Métriques business (Odoo automatique)", 3)
add_table(
    headers=["Métrique", "Cible"],
    rows=[
        ["Commandes générées par les destinataires (9 juin → 9 juillet)", "16-18"],
        ["CA additionnel",                                                  "+9 K€"],
        ["Réactivations Revendeur dormant",                                 "13-15 sur 105 (~13%)"],
        ["Réactivations Horeca dormant",                                    "3 sur 11 (sur petit volume)"],
    ],
    col_widths_cm=[16, 10]
)
add_hr()

# ----- Workflow récap -----
add_heading("Workflow récap (qui fait quoi)", 2)
p = doc.add_paragraph()
p.style = doc.styles["Normal"]
workflow_lines = [
    "NICOLAS  → segmente Odoo, tag, exporte CSV (email + nom uniquement)",
    "            ↓",
    "STEPHAN  → importe CSV Mailchimp, crée 2 audiences, designe 2 templates",
    "            ↓",
    "            envoie les 2 mailings le 9 juin",
    "            ↓",
    "CLIENT   → reçoit email, passe commande en mentionnant l'offre (sans code)",
    "            ↓",
    "VANESSA  → reçoit la commande, applique l'offre sur le devis Odoo",
    "            (3+1 ou -10% selon segment) — pas de code à saisir",
    "            ↓",
    "JÉRÔME   → appelle top 20 Revendeur + top 5 Horeca dormants à J+10",
    "            ↓",
    "NICOLAS  → J+30 reporting : croisement listes Odoo → conversions, CA, ROI",
    "            → ajuste la campagne juillet (Collection Glacée) en fonction",
]
for line in workflow_lines:
    r = p.add_run(line + "\n")
    r.font.name = "Consolas"
    r.font.size = Pt(9)
add_hr()

# ----- Points d'attention -----
add_heading("Points d'attention", 2)
attentions = [
    ("Pas de chevauchement",   "un même partner = 1 seul tag (dedup déjà faite). Aucun client ne recevra 2 mailings différents."),
    ("Bounces probables",      "sur 116 dormants, attendre 5-10 bounces (emails obsolètes). Vanessa nettoiera la liste Odoo en parallèle."),
    ("Anti-spam",              "planifie les 2 envois espacés : 9h00 → Revendeur dormant (105), 9h30 → Horeca dormant (11)."),
    ("A/B testing",            "si Mailchimp Pro le permet, tester 2 objets sur 10-20% de la base Revendeur dormant — l'objet gagnant part au reste."),
    ("Pas d'envoi aux Horeca actifs", "les 178 cafés/restos qui commandent activement ne reçoivent RIEN cette campagne (décision Nicolas : éviter la cannibalisation). Ils restent en base, on les sollicitera plus tard avec une mécanique adaptée."),
    ("Limite tracking sans code", "si un client transfère le mail à un copain qui n'est pas dans la liste, on ne le saura pas. Acceptable car volume marginal sur 116 dormants."),
]
for i, (title, desc) in enumerate(attentions, start=1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    p.add_run(title + " : ").bold = True
    p.add_run(desc)
add_hr()

# ----- Prochaine étape -----
add_heading("Prochaine étape pour toi", 2)
add_para("1. Lis ce brief, dis-moi si l'angle / le ton / le calendrier te conviennent")
add_para("2. Si OK, on bloque un point de 30 min cette semaine pour aligner sur les visuels")
add_para("3. Tu peux déjà commencer à brainstormer les objets email et les accroches")
add_hr()

# ----- Fichiers -----
add_heading("Fichiers de référence", 2)
add_bullet("CSV Mailchimp : data/mailchimp_juin_2026.csv (112 lignes, sans code promo)")
add_bullet("Roadmap complète : Roadmap_B2B_2026-2027.pdf")
add_bullet("Proposition offre juin : Offre_Juin_2026_Proposition.pdf")
add_bullet("Liste opérationnelle Vanessa : Liste_Vanessa_Juin_2026.xlsx")

doc.save(OUT)
print(f"DOCX genere : {OUT}")
print(f"Taille : {OUT.stat().st_size / 1024:.1f} Ko")
